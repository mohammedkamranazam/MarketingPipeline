"""
Tests for parser_service and csv_parser_service.

Acceptance criteria tested:
- PDF/DOCX/TXT parse_document returns ParsedPage list.
- Unsupported mime_type raises ValueError.
- CSV/XLSX parse_tabular returns TabularParseResult with headers and rows.
- Unsupported tabular mime_type raises ValueError.
- Empty CSV returns empty headers and rows.
- XLSX None cells coerced to "".
"""

import csv as csv_module
import io

import pytest

from core.services.csv_parser_service import parse_tabular
from core.services.parser_service import parse_document

# ---------------------------------------------------------------------------
# Parser service
# ---------------------------------------------------------------------------

def test_parse_txt_single_page():
    data = b"Hello world"
    pages = parse_document(data, "text/plain")
    assert len(pages) == 1
    assert pages[0].page_number == 1
    assert pages[0].raw_text == "Hello world"


def test_parse_txt_utf8_decode():
    data = "Héllo wörld".encode()
    pages = parse_document(data, "text/plain")
    assert "H" in pages[0].raw_text


def test_parse_unsupported_mime_raises():
    with pytest.raises(ValueError, match="Unsupported mime_type"):
        parse_document(b"data", "application/zip")


def test_parse_pdf_single_page():
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pages = parse_document(buf.getvalue(), "application/pdf")
    assert len(pages) == 1
    assert pages[0].page_number == 1
    assert isinstance(pages[0].raw_text, str)


def test_parse_pdf_two_pages():
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pages = parse_document(buf.getvalue(), "application/pdf")
    assert len(pages) == 2
    assert pages[1].page_number == 2


def test_parse_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    pages = parse_document(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert len(pages) == 1
    assert "First paragraph" in pages[0].raw_text
    assert "Second paragraph" in pages[0].raw_text


# ---------------------------------------------------------------------------
# CSV parser service
# ---------------------------------------------------------------------------

def _make_csv(rows: list[dict]) -> bytes:
    if not rows:
        return b""
    buf = io.StringIO()
    writer = csv_module.DictWriter(buf, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def test_parse_csv_basic():
    data = _make_csv([
        {"first_name": "Alice", "company": "Acme"},
        {"first_name": "Bob", "company": "Beta"},
    ])
    result = parse_tabular(data, "text/csv")
    assert result.headers == ["first_name", "company"]
    assert len(result.rows) == 2
    assert result.rows[0]["first_name"] == "Alice"


def test_parse_csv_empty_returns_empty():
    result = parse_tabular(b"", "text/csv")
    assert result.headers == []
    assert result.rows == []


def test_parse_tabular_unsupported_mime_raises():
    with pytest.raises(ValueError, match="Unsupported tabular mime_type"):
        parse_tabular(b"data", "application/pdf")


def test_parse_xlsx_basic():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["first_name", "company"])
    ws.append(["Carol", "Gamma"])
    buf = io.BytesIO()
    wb.save(buf)
    result = parse_tabular(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert result.headers == ["first_name", "company"]
    assert result.rows[0]["company"] == "Gamma"


def test_parse_xlsx_none_cells_become_empty():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["first_name", "company"])
    ws.append([None, "Delta"])
    buf = io.BytesIO()
    wb.save(buf)
    result = parse_tabular(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert result.rows[0]["first_name"] == ""
